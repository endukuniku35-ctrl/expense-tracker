import os
import sys
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def set_cell_background(cell, fill_color):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_color}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="CCCCCC", sz="4", val="single"):
    tblPr = table._element.xpath('w:tblPr')
    if tblPr:
        borders = parse_xml(f'''
            <w:tblBorders {nsdecls("w")}>
                <w:top w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:bottom w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideH w:val="{val}" w:sz="{sz}" w:space="0" w:color="{color}"/>
                <w:insideV w:val="none"/>
                <w:left w:val="none"/>
                <w:right w:val="none"/>
            </w:tblBorders>
        ''')
        tblPr[0].append(borders)

def build_docx(filename):
    doc = Document()
    
    # Page Margins (1 inch all sides)
    sections = doc.sections
    for s in sections:
        s.top_margin = Inches(1)
        s.bottom_margin = Inches(1)
        s.left_margin = Inches(1)
        s.right_margin = Inches(1)

    # Styling Colors
    COLOR_PRIMARY = RGBColor(26, 115, 232)     # #1A73E8
    COLOR_SECONDARY = RGBColor(13, 71, 161)   # #0D47A1
    COLOR_DARK = RGBColor(30, 41, 59)         # #1E293B
    COLOR_MUTED = RGBColor(100, 116, 139)     # #64748B

    # Document Header / Banner Table
    banner_table = doc.add_table(rows=1, cols=1)
    banner_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    banner_cell = banner_table.cell(0, 0)
    set_cell_background(banner_cell, "0F172A")
    set_cell_margins(banner_cell, top=200, bottom=200, left=250, right=250)

    p_title = banner_cell.paragraphs[0]
    p_title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_title = p_title.add_run("🍛 CurryTracker")
    run_title.font.name = 'Arial'
    run_title.font.size = Pt(24)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(255, 255, 255)

    p_sub = banner_cell.add_paragraph()
    run_sub = p_sub.add_run("System Architecture, Multi-Tenant Design & Complete Implementation Guide")
    run_sub.font.name = 'Arial'
    run_sub.font.size = Pt(12)
    run_sub.font.color.rgb = RGBColor(147, 197, 253)

    p_meta = banner_cell.add_paragraph()
    run_meta = p_meta.add_run("Live Web App: https://expense-tracker-77br.onrender.com | Version: 49.0 | Date: July 2026")
    run_meta.font.name = 'Arial'
    run_meta.font.size = Pt(9.5)
    run_meta.font.color.rgb = RGBColor(203, 213, 225)

    doc.add_paragraph() # Spacer

    # Helper function for Headings
    def add_h1(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(6)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return h

    def add_h2(text):
        h = doc.add_paragraph()
        h.paragraph_format.space_before = Pt(12)
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.keep_with_next = True
        run = h.add_run(text)
        run.font.name = 'Arial'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return h

    def add_body(text, bold_prefix=None):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_b = p.add_run(bold_prefix)
            run_b.font.name = 'Calibri'
            run_b.font.size = Pt(11)
            run_b.font.bold = True
            run_b.font.color.rgb = COLOR_DARK
        run_t = p.add_run(text)
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(11)
        run_t.font.color.rgb = COLOR_DARK
        return p

    def add_bullet(bold_txt, body_txt):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.15
        run_b = p.add_run(bold_txt + ": ")
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10.5)
        run_b.font.bold = True
        run_b.font.color.rgb = COLOR_DARK
        run_t = p.add_run(body_txt)
        run_t.font.name = 'Calibri'
        run_t.font.size = Pt(10.5)
        run_t.font.color.rgb = COLOR_DARK
        return p

    def add_code(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(8)
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = COLOR_DARK
        return p

    # --- SECTION 1: EXECUTIVE SUMMARY ---
    add_h1("1. Executive Summary")
    add_body(
        "CurryTracker is a production-grade, multi-tenant web application designed for managing shared room expenses, "
        "daily curry/meal splits, UPI payment settlements, and roommate communication. Built with a lightweight Node.js/Express "
        "backend and an interactive Vanilla JavaScript Single-Page Application (SPA) frontend, CurryTracker provides zero-latency "
        "expense tracking, dynamic per-meal cost calculation, 2-second mobile push notifications, dynamic custom UPI QR code generation, "
        "and 1-tap Telegram group report broadcasting."
    )

    # Table 1: Key Feature Summary
    t1 = doc.add_table(rows=6, cols=2)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t1)
    
    headers = ["Feature Module", "Implementation & Capabilities"]
    for i, h in enumerate(headers):
        cell = t1.cell(0, i)
        set_cell_background(cell, "EFF6FF")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_SECONDARY

    rows_data = [
        ("Multi-Tenant RBAC", "3-tier role hierarchy: Main Super Admin, Secondary Group Admins, and Standard Members."),
        ("Dynamic Split Engine", "Dynamic per-meal cost calculation based on participating member array (splitBetween)."),
        ("Dual Persistence Engine", "Embedded SQLite3 database synchronized with JSON data files with auto-healing fallback."),
        ("Mobile Web Push Engine", "2000ms polling push engine delivering status bar alerts with sound and vibration patterns."),
        ("UPI & Telegram Integration", "Dynamic custom amount UPI QR generator (upiqr.in API) & 1-tap Telegram markdown share.")
    ]

    for row_idx, (feat, desc) in enumerate(rows_data, start=1):
        cell_a = t1.cell(row_idx, 0)
        cell_b = t1.cell(row_idx, 1)
        
        p_a = cell_a.paragraphs[0]
        run_a = p_a.add_run(feat)
        run_a.font.name = 'Calibri'
        run_a.font.bold = True
        run_a.font.size = Pt(10)

        p_b = cell_b.paragraphs[0]
        run_b = p_b.add_run(desc)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(10)

    doc.add_paragraph()

    # --- SECTION 2: SYSTEM ARCHITECTURE & TECH STACK ---
    add_h1("2. System Architecture & Technology Stack")
    add_body("The application follows a decoupled SPA architecture operating over a RESTful API backend:")

    # Table 2: Tech Stack
    t2 = doc.add_table(rows=9, cols=3)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t2)

    stack_headers = ["Layer", "Technologies Used", "Core Purpose"]
    for i, h in enumerate(stack_headers):
        cell = t2.cell(0, i)
        set_cell_background(cell, "EFF6FF")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_SECONDARY

    stack_data = [
        ("Frontend Core", "HTML5, Vanilla JS (ES6+ SPA)", "60fps high performance UI, modular view routing without heavy framework bundles."),
        ("UI & Aesthetics", "Glassmorphism, CSS Tokens, Bootstrap 5", "Vibrant dark/light hybrid design system, responsive modals, animated stat cards."),
        ("Analytics & Charts", "Chart.js v4", "Interactive doughnut pie charts, monthly line trends, and weekly bar charts."),
        ("Backend Framework", "Node.js, Express.js", "RESTful API architecture, route handling, express-session cookie management."),
        ("Security & Auth", "Bcrypt.js, Express-Session, RBAC", "Salted password hashing (factor 10), session isolation, role enforcement middleware."),
        ("Database Layer", "SQLite3 + Synchronized JSON Files", "Dual-storage fallback engine ensuring zero data loss across cloud deployments."),
        ("Real-Time Engine", "Mobile Push API + 2000ms Polling", "Instant status bar push notifications with sound, vibration, and custom icons."),
        ("Integrations", "upiqr.in API, Telegram Web API", "Dynamic custom amount UPI QR generation & 1-tap Telegram group report sharing.")
    ]

    for row_idx, (layer, tech, purp) in enumerate(stack_data, start=1):
        for col_idx, text in enumerate([layer, tech, purp]):
            cell = t2.cell(row_idx, col_idx)
            p = cell.paragraphs[0]
            run = p.add_run(text)
            run.font.name = 'Calibri'
            run.font.size = Pt(9.5)
            if col_idx == 0:
                run.font.bold = True

    doc.add_paragraph()

    # --- SECTION 3: MULTI-TENANT ROLE HIERARCHY ---
    add_h1("3. Multi-Tenant Role-Based Access Control (RBAC)")
    add_body("CurryTracker implements a 3-tier security model enforcing explicit permissions across backend APIs and UI components:")

    add_bullet("👑 Main Super Admin (super_admin)", "Kandukuri Jagan (User ID: 192472374). Has systemwide master access across all room groups, total platform expense analytics, ability to create Secondary Group Admins, view password credentials directory with unhide toggles (👁️), and remove/delete (🗑️) any account.")
    add_bullet("🛡️ Secondary Group Admin (admin)", "e.g., anudeep (User ID: 192472066). Has an isolated dashboard dedicated exclusively to their specific room/group members. Can add regular Member accounts and record room expenses & cash settlements. Hidden from credentials directory & super admin controls.")
    add_bullet("👤 Standard Member (member)", "e.g., Sagar, Prathap, Bharath, Charan, Ganesh, a. Has an isolated view restricted ONLY to expenses where they paid or participated in the split. Can view their personal balance sheet ('🔴 Owes', '💰 Receives', '✅ Settled') and record payment settlements.")

    doc.add_paragraph()

    # --- SECTION 4: CORE MATHEMATICAL ALGORITHMS & LOGIC ---
    add_h1("4. Core Mathematical Algorithms & System Logic")
    
    add_h2("4.1 Dynamic Per-Meal Split Formula")
    add_body("CurryTracker dynamically calculates meal split shares using the following mathematical model:")

    split_code = (
        "Share per Person (Meal i) = Amount(i) / Count(splitBetween_i)\n"
        "Total Paid (Member m) = Σ Amount(i) where paidBy = m\n"
        "Total Share (Member m) = Σ (Amount(i) / Count(splitBetween_i)) where m ∈ splitBetween_i\n"
        "Net Balance (Member m) = (Total Paid + Settlements Out) - (Total Share + Settlements In)"
    )
    add_code(split_code)

    add_h2("4.2 Persistent Dual-Storage & Auto-Healing Engine")
    add_body("To prevent data loss across cloud deployments (such as Render ephemeral container restarts):")
    add_bullet("SQLite3 Primary Engine", "Executes SQL transactions for high-speed queries and data indexing.")
    add_bullet("JSON Backup Persistence", "State is synchronized to formatted JSON files (data/users.json, data/expenses.json, data/settlements.json).")
    add_bullet("Auto-Healing Initialization", "On server startup, initDatabase() inspects, merges, and self-heals missing accounts or sample data.")

    doc.add_paragraph()

    # --- SECTION 5: CODEBASE STRUCTURE ---
    add_h1("5. Codebase Structure & File Taxonomy")
    add_body("The project is structured logically into backend routes, middleware, database modules, and frontend controllers:")

    file_tree = (
        "expense-tracker/\n"
        "├── server.js               # Entry point: Express app, session setup, route mounting\n"
        "├── database.js             # SQLite3 connection, JSON sync, schema init & seed data\n"
        "├── middleware/auth.js      # Auth guards (requireAuth, requireAdmin, requireSuperAdmin)\n"
        "├── routes/\n"
        "│   ├── auth.js             # User login, session check, logout\n"
        "│   ├── members.js          # Member creation, Group Admin creation, credentials directory, user removal\n"
        "│   ├── expenses.js         # Expense CRUD operations & group filtering\n"
        "│   ├── balance.js          # Balance calculation engine & settlement recording\n"
        "│   ├── dashboard.js        # Aggregated stats & Chart.js analytics API\n"
        "│   ├── messages.js         # Roommate chat & admin broadcast endpoints\n"
        "│   └── reports.js          # Monthly & member export reports (CSV/Excel/PDF)\n"
        "└── public/\n"
        "    ├── dashboard.html      # SPA HTML shell with view containers & Bootstrap modals\n"
        "    ├── css/style.css       # Core Design System, CSS Variables, Glassmorphism styles\n"
        "    └── js/\n"
        "        ├── app.js          # SPA Router, Session state, Navigation guards\n"
        "        ├── dashboard.js    # Dashboard renderer, counter animations, Chart.js instances\n"
        "        ├── members.js      # Members cards rendering, login directory, Group Admin modal\n"
        "        ├── expenses.js     # Expense history table, Add/Edit modal handlers\n"
        "        ├── payments.js     # Payment balance sheet, settlements table, custom amount UPI QR generator\n"
        "        ├── messages.js     # Roommate chat & Broadcast banner UI controllers\n"
        "        └── notifications.js # 2-second ultra-fast mobile push notification manager"
    )
    add_code(file_tree)

    doc.add_paragraph()

    # --- SECTION 6: KEY FEATURES & INTEGRATIONS ---
    add_h1("6. Key Features & External Integrations")
    add_bullet("Dynamic Custom Amount UPI QR Generator", "Powered by the upiqr.in API. Allows entering custom settlement amounts or tapping '🔴 Owes'. Generates instant scannable UPI QR codes targeting verified VPA 8367047947@ybl with deep links to PhonePe, Google Pay, and Paytm.")
    add_bullet("2-Second Mobile Push Notification Engine", "2000ms browser notification synchronizer. Native push notifications pop up directly in the Android/iOS status bar with custom sound, vibration patterns ([500, 200, 500, 200, 500]), and app crown icons.")
    add_bullet("1-Tap Telegram Group Report Share", "Share on Telegram buttons across Dashboard, Reports, and Payments views. Automatically compiles itemized Markdown reports and opens Telegram Web or Desktop app via t.me/share to broadcast to room groups.")

    doc.add_paragraph()

    # --- SECTION 7: DEPLOYMENT & OPERATIONAL DETAILS ---
    add_h1("7. Operational Credentials & Deployment Specifications")
    
    t3 = doc.add_table(rows=7, cols=2)
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(t3)

    ops_headers = ["Parameter", "Configuration Value / Credentials"]
    for i, h in enumerate(ops_headers):
        cell = t3.cell(0, i)
        set_cell_background(cell, "EFF6FF")
        p = cell.paragraphs[0]
        run = p.add_run(h)
        run.font.name = 'Arial'
        run.font.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = COLOR_SECONDARY

    ops_rows = [
        ("Live Hosted Application", "https://expense-tracker-77br.onrender.com"),
        ("GitHub Repository", "https://github.com/endukuniku35-ctrl/expense-tracker.git"),
        ("Main Super Admin Account", "User ID: 192472374 | Password: kandukurijagan@14062020 (Kandukuri Jagan)"),
        ("Secondary Group Admin Account", "User ID: anudeep | Password: anudeep (Group Admin Flat 301)"),
        ("Standard Member Accounts", "User IDs: 192472343 (Sagar), 192411184 (Prathap), 192411185 (Bharath), 192412348 (Charan), Ganesh (Ganesh), a (Member a)"),
        ("Default Payment VPA", "UPI ID: 8367047947@ybl")
    ]

    for row_idx, (param, val) in enumerate(ops_rows, start=1):
        cell_a = t3.cell(row_idx, 0)
        cell_b = t3.cell(row_idx, 1)
        
        p_a = cell_a.paragraphs[0]
        run_a = p_a.add_run(param)
        run_a.font.name = 'Calibri'
        run_a.font.bold = True
        run_a.font.size = Pt(9.5)

        p_b = cell_b.paragraphs[0]
        run_b = p_b.add_run(val)
        run_b.font.name = 'Calibri'
        run_b.font.size = Pt(9.5)

    doc.save(filename)
    print(f"Successfully generated DOCX at: {filename}")

if __name__ == '__main__':
    out_dir = os.path.dirname(os.path.abspath(__file__))
    target_docx = os.path.join(out_dir, "CurryTracker_Architecture_and_Creation_Guide.docx")
    build_docx(target_docx)
